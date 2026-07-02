from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(102, 102, 102)

STRUCTURE_HELP = """\
Recommended structure:

  my_document/
    sections/
      00_cover.md
      01_summary.md
      02_architecture.md
    assets/
      context_diagram.png
      data_flow.png
    dist/
      document.docx

Recommended usage:

  python tools/md_to_docx_cli.py my_document/sections my_document/dist/document.docx

You can also pass a single Markdown file:

  python tools/md_to_docx_cli.py README.md dist/README.docx

Supported Markdown:

  # Heading 1
  ## Heading 2
  ### Heading 3

  Normal paragraphs separated by blank lines.

  - Simple list item
  - Another item

  | Column A | Column B |
  | --- | --- |
  | Value | Value |

  ![Figure 1. Caption](../assets/diagram.png){width=16}

  {{pagebreak}}

Notes:

  - If you pass a folder, .md files are processed alphabetically.
  - Use numeric prefixes to control order: 00_, 01_, 02_...
  - Image paths are resolved from the Markdown file where they appear.
  - Image width is expressed in centimeters and capped at 17 cm.
"""


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
    return text


def set_run(
    run,
    size: float = 9.4,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def parse_shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.4)
    for name, size in [("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 10.8)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE


def add_paragraph(
    doc: Document,
    text: str,
    size: float = 9.4,
    bold: bool = False,
    italic: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(clean_inline(text))
    set_run(run, size=size, bold=bold, italic=italic)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    separator = r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?"
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(separator, stripped):
            continue
        rows.append([clean_inline(cell) for cell in stripped.strip("|").split("|")])
    max_cols = max((len(row) for row in rows), default=0)
    return [row + [""] * (max_cols - len(row)) for row in rows]


def add_table(doc: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run(
                run,
                size=8,
                bold=(r_idx == 0),
                color=RGBColor(255, 255, 255) if r_idx == 0 else None,
            )
            if r_idx == 0:
                cell._tc.get_or_add_tcPr().append(parse_shading("1F4E79"))
    doc.add_paragraph()


def add_image(doc: Document, line: str, md_path: Path) -> None:
    match = re.match(r"!\[(.*?)\]\((.*?)\)(?:\{width=([0-9.]+)\})?", line.strip())
    if not match:
        add_paragraph(doc, line)
        return

    caption, raw_path, raw_width = match.groups()
    img_path = (md_path.parent / raw_path).resolve()
    if not img_path.exists():
        add_paragraph(doc, f"[Image not found: {raw_path}]", italic=True)
        return

    width_cm = min(float(raw_width or 16), 17.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        set_run(cap_run, size=8, color=GRAY)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(clean_inline(text))
    set_run(run, size=9, italic=True, color=GRAY)


def md_to_docx(doc: Document, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    paragraph: list[str] = []
    table_lines: list[str] = []
    bullets: list[str] = []
    in_comment = False

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            add_paragraph(doc, " ".join(paragraph))
            paragraph = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    def flush_bullets() -> None:
        nonlocal bullets
        for item in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean_inline(item))
            set_run(run, size=9.4)
        bullets = []

    for raw in lines:
        line = raw.rstrip().lstrip("\ufeff")
        stripped = line.strip()

        if in_comment:
            if "-->" in line:
                in_comment = False
            continue
        if stripped.startswith("<!--"):
            in_comment = "-->" not in line
            continue
        if stripped == "{{pagebreak}}":
            flush_paragraph()
            flush_table()
            flush_bullets()
            if doc.paragraphs:
                doc.add_page_break()
            continue
        if not stripped:
            flush_paragraph()
            flush_table()
            flush_bullets()
            continue
        if line.startswith("|"):
            flush_paragraph()
            flush_bullets()
            table_lines.append(line)
            continue
        if line.startswith("- "):
            flush_paragraph()
            flush_table()
            bullets.append(line[2:].strip())
            continue
        if line.startswith(">"):
            flush_paragraph()
            flush_table()
            flush_bullets()
            add_quote(doc, line.lstrip("> ").strip())
            continue
        if line.startswith("!"):
            flush_paragraph()
            flush_table()
            flush_bullets()
            add_image(doc, line, md_path)
            continue
        if line.startswith("#"):
            flush_paragraph()
            flush_table()
            flush_bullets()
            level = min(len(line) - len(line.lstrip("#")), 3)
            text = line[level:].strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            run = p.add_run(clean_inline(text))
            set_run(run, size={1: 16, 2: 12.5, 3: 10.8}[level], bold=True, color=BLUE)
            continue

        paragraph.append(stripped)

    flush_paragraph()
    flush_table()
    flush_bullets()


def collect_markdown_files(input_path: Path, pattern: str, recursive: bool) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    if not input_path.is_dir():
        raise FileNotFoundError(f"Input path does not exist: {input_path}")

    files = input_path.rglob(pattern) if recursive else input_path.glob(pattern)
    return sorted(path for path in files if path.is_file())


def build_docx(input_path: Path, output_path: Path, pattern: str, recursive: bool) -> Path:
    md_files = collect_markdown_files(input_path, pattern, recursive)
    if not md_files:
        raise FileNotFoundError(
            f"No Markdown files matching pattern {pattern!r} were found in {input_path}"
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)

    for index, md_path in enumerate(md_files):
        if index > 0:
            doc.add_page_break()
        md_to_docx(doc, md_path)

    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert one or more Markdown files to a Word .docx document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=STRUCTURE_HELP,
    )
    parser.add_argument(
        "input",
        nargs="?",
        type=Path,
        help="A .md file or a folder with .md files. Folders are processed by filename.",
    )
    parser.add_argument(
        "output",
        nargs="?",
        type=Path,
        help="Output .docx path. Example: dist/document.docx.",
    )
    parser.add_argument(
        "--pattern",
        default="*.md",
        help="Markdown file pattern when input is a folder. Default: *.md",
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Also search for Markdown files in subfolders.",
    )
    parser.add_argument(
        "--print-structure",
        action="store_true",
        help="Print the recommended structure and exit.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.print_structure:
        print(STRUCTURE_HELP)
        return

    if not args.input or not args.output:
        raise SystemExit("Usage: python tools/md_to_docx_cli.py <input.md|folder> <output.docx>")

    output_path = build_docx(args.input, args.output, args.pattern, args.recursive)
    print(output_path)


if __name__ == "__main__":
    main()
